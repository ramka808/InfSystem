import docx
from docx import Document
from docx.shared import Inches
import sys

arg = sys.argv[:]

arg2 = []
arg3 = []
#Разделить входной текст на массив массивов(которые состоят из строки) затем эти строки надо разделить на отдельные объекты
for i in range(len(arg)):

    arg2.append(arg[i].split("'"))


#разделение каждого элемента массива по запятым, после которых не следует пробел, это запятые которые отделяют
# массивы внутри массива. Для этого запятые с пробелами заменяются на символ доалара
for i in arg2:
    i[0] = i[0].replace(', ', '$')
    arg3.append(i[0].split(','))


#обратная замена доллара на запятую с пробелом
for i in range(len(arg3)):
    for j in range(len(arg3[i])):
        arg3[i][j] = arg3[i][j].replace('$',', ')




datd2 = [
    ['Группа 4297 Направление 09.04.02 Информационные системы и технологии'],
    [1, 'Гаптуллин Марат Рафисович',
     'Панченко Оксана Владимировна, канд. техн. наук, доцент, доцент кафедры «Интеллектуальных систем и управления информационными ресурсами» КНИТУ',
     'Шлеймович М.П., канд.техн.наук, зав. каф. АСОИУ']
    , [2, 'Гатиятуллин Вадим Айратович',
       'Панченко Оксана Владимировна, канд. техн. наук, доцент, доцент кафедры «Интеллектуальных систем и управления информационными ресурсами» КНИТУ',
       'Сытник А.С., канд.техн.наук, доцент каф. АСОИУ'],
    [3, 'Замалетдинов Ильнур Агзамович',
     'Голицына И.Н., канд.физ.-мат.наук, доцент кафедры инжиниринга ПО Высшей школы ИТИС К(П)ФУ',
     'Эминов Ф.И., канд.техн.наук, доцент каф. АСОИУ'],
[1, 'Замалетдинов Ильнур Агзамович',
                                                                                'Голицына И.Н., канд.физ.-мат.наук, доцент кафедры инжиниринга ПО Высшей школы ИТИС К(П)ФУ',
                                                                                'Эминов Ф.И., канд.техн.наук, доцент каф. АСОИУ'],
    ['Группа 4296 Направление 09.04.02 Информационные системы и технологии ']
]

docInput = docx.Document()
docInput.add_paragraph(str(arg3))
docInput.save('inputAr.docx')
doc = docx.Document('input.docx')

table = doc.tables[1]
table.style = 'Table Grid'

# cell1 = table.cell(0, 0)
# cell2 = table.cell(0, 1)
# cell3 = table.cell(0, 2)
# cell4 = table.cell(0, 3)
# cell1.merge(cell2)
# cell1.merge(cell3)
# cell1.merge(cell4)
for row_data in arg3:
    new_row = table.add_row()
    for i, cell_data in enumerate(row_data):
        new_cell = new_row.cells[i]
        new_cell.text = str(cell_data)
for i in range(0, len(arg3)):
    if len(arg3[i]) == 1:
        cell1 = table.cell(i+1, 0)
        cell2 = table.cell(i+1, 1)
        cell3 = table.cell(i+1, 2)
        cell4 = table.cell(i+1, 3)
        cell1.merge(cell2)
        cell1.merge(cell3)
        cell1.merge(cell4)
# # Добавляем строки из второй таблицы в первую таблицу
# for row in table2.rows:
#     new_row = table.add_row()
#     for i, cell in enumerate(row.cells):
#         new_cell = new_row.cells[i]
#         new_cell.text = cell.text
# for rows in range(len(datd2)):
# # добавляем новую строку и заполняем ее данными из списка new_row_data
# new_row_data = ['Новая строка', 'Данные 1', 'Данные 2']
# new_row = table.add_row()
# for i, cell_data in enumerate(new_row_data):
#     new_cell = new_row.cells[i]
#     new_cell.text = cell_data
# сохраняем изменения в документе
doc.save('output.docx')

# Title
# Heading 1
# Heading 2
# Heading 3
# Heading 4
# Heading 5
# Heading 6
# Strong
# List Paragraph
# footnote text

